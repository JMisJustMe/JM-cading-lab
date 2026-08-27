from pathlib import Path
from datetime import datetime
from hashlib import sha256

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether

OUT = Path(__file__).resolve().parent / "dist"
OUT.mkdir(parents=True, exist_ok=True)
BASENAME = "JM_UNIVERSAL_HOSTED_BODY_Architecture_Proof_Gate_Closure_v1_0"
DOCX = OUT / f"{BASENAME}.docx"
PDF = OUT / f"{BASENAME}.pdf"

TITLE = "JM UNIVERSAL HOSTED-BODY ARCHITECTURE"
SUBTITLE = "Identity-Preserving Embedded Function Across Hosts — Architecture & Proof-Gate Closure v1.0"
DATE = "28 August 2026 · Europe/London"
CANONICAL_ROUTE = "/Zionfolder/JM Estate/Runtime & Host Architecture/Universal Hosted-Body/Architecture and Proof-Gate Closure v1.0/"

BODY_ID = "jm.untitled-field-branch/v0.9"
BODY_SHA = "7391dd5bc1c4ff1565d70b69354cfdd79a121f8e6f6a1d671d75b57d463ee7ea"
PLAYABLE_IR = "jm.gamecore.playable-ir/v0.3"
ONEBODY_ABI = "jm.onebody-abi/v0.1"
PROOF_BRANCH = "universal-hosted-body-exact-identity-proof-v0.1"
PROOF_PR = "JMisJustMe/JM-cading-lab PR #166"
PROOF_HEAD = "5e9f958ee159d24c795ce1389129436bff64ccbe"

CORE = [
    ("FUNCTION", "What the software does."),
    ("BODY", "The identity-bearing game, tool, app or interface."),
    ("HOST", "The environment that contains or executes the body."),
    ("ADAPTER", "The translation layer between body requirements and host capabilities."),
    ("MANIFESTATION", "How the body appears there: inline, panel, plugin, overlay, WebView, standalone surface, and so on."),
]

SECTIONS = [
    ("1. Canonical determination", [
        "The original ‘inline thing’ is not a special identity in itself. In ordinary technical terms it is an embedded interactive interface or mini-app operating inside another host surface. ‘Inline’ describes one manifestation position; it does not name the underlying software body.",
        "The broader principle is host-agnostic: a functioning software body may operate within another functioning host without becoming that host or necessarily surrendering its own identity, provided the host supplies — directly or through adaptation — the capabilities the body requires.",
        "The universal relationship is therefore BODY ↔ ADAPTER ↔ HOST → MANIFESTATION. ChatGPT is one host. Inline is one manifestation. A widget is one possible implementation form. JM INLINE CONTACT governs the inline/contact branch rather than owning the universal principle.",
    ]),
    ("2. Core invariant and identity law", [
        "BODY ≠ HOST ≠ ADAPTER ≠ MANIFESTATION.",
        "A change of host or presentation does not, by itself, create a new software identity. Identity is preserved when the defining function and authoritative core remain materially the same and host-specific differences are carried primarily by adapters, bindings or carriers.",
        "Host-agnostic principle ≠ identical implementation everywhere. Different hosts expose different runtimes, APIs, input systems, storage models, permission boundaries and rendering surfaces; adapters reconcile those differences without needlessly rewriting the body core.",
    ]),
    ("3. Capability model", [
        "Every body declares what it needs. Every host declares what it can provide. A matcher classifies each requirement as DIRECT, ADAPTED/FALLBACK, OPTIONAL-MISSING, ENHANCEMENT-MISSING or BLOCKING-MISSING.",
        "Compatibility is therefore richer than a binary ‘works / does not work’. A body may retain full defining function while substituting pointer for touch, session state for persistent state, or omitting non-essential audio/haptics. Any degradation must remain explicit.",
        "OPENED ≠ FUNCTIONED ≠ FULLY FUNCTIONED. A manifestation earns stronger status only when direct contact produces the expected state change and consequence under the claimed host conditions.",
    ]),
    ("4. Adapter contract", [
        "The shared adapter surface is conceptual rather than host-specific: discoverCapabilities; attach/mount; detach/unmount; map input; map output; render; load/save state; request permissions; emit/receive events; report degradation/failures; run function proof; run identity proof.",
        "Adapters may themselves be composed from reusable organs such as runtime, input, rendering, state, storage, network, permission and event bridges. More than one host can share an adapter when they expose a sufficiently common interface.",
    ]),
    ("5. Existing Estate federation — recovered, not rebuilt", [
        "The universal model does not require a new game engine or general runtime. Existing organs already occupy the needed roles: OneBody IR / jm.onebody-abi v0.1 for host-neutral identity carriage; JM GameCore playable IR v0.3 for playable semantics; Body House / JM CodeHouse / MultiHub for body and capability registry; GlyphForge for adapter/capability work; PolyBridge Stack for cross-body bridges; JM32-1DA Runtime and Cross-Device Runtime Adapter for routing; JM Visual Interaction Runtime for input→state→visible consequence; OneBody Delivery for multi-channel carriage; TraceBox / Code Ding / Receipt Shape for proof.",
        "JM ChatGPT Host Adapter v0.1.1/v0.1.2 is the thin ChatGPT/MCP host organ. Its governing rule is CARRIER_NOT_SOURCE_AUTHORITY with mergeForbidden=true. The host carries the body; it does not become source authority merely by rendering it.",
    ]),
    ("6. Recovered evidence before this closure", [
        "PAIRBREAK v0.7.2 already supplied useful cross-host evidence: the exact frozen body passed a broad Chromium mobile-profile mechanical suite and was physically contacted on Android. That supports host continuity, but it is not the exact same-body ChatGPT proof used for the final join below.",
        f"Untitled Field Branch v0.9 is the stronger exact-identity reference. Its frozen keeper SHA-256 is {BODY_SHA}. Records show 27/27 phone-path QA, desktop pass, gesture-contact pass, zero page errors and owner-device acceptance.",
        "The historical ChatGPT adapter pointer names that exact frozen file, pins the same SHA, declares the same playable IR and OneBody ABI, forbids mutation and declares ChatGPT to be carrier rather than source authority. MCP runtime/tool/resource transport was proved, but the adapter receipt correctly stated that direct ChatGPT contact of this exact donor was not proved there.",
        "Separately, JM INLINE CONTACT Third Ding proved the host class itself: a newly produced embedded body inside the ChatGPT conversation visibly changed READY / 0 → DING 1/1 → DING 5/5 under direct user contact. That proves genuine interactive inline contact, but it must not be silently substituted for exact Untitled Field donor contact.",
    ]),
    ("7. Exact empirical join — deliberately narrow", [
        "The one remaining empirical join is: run the exact frozen Untitled Field Branch v0.9 body through the existing JM ChatGPT Host Adapter, make direct user contact with that exact body inline, observe visible state consequence, and record the same pinned identity in the resulting receipt.",
        "Only that event earns EXACT-BODY CROSS-HOST DING: PROVED. Universal arbitrary-host portability is a larger claim and remains uncrowned even after that join.",
    ]),
    ("8. 28 August 2026 proof-rail implementation", [
        f"A non-destructive descendant branch was created from the historical adapter line: {PROOF_BRANCH}. The historical branch and frozen donor remain unchanged.",
        f"A draft proof PR was opened: {PROOF_PR}. The closure is pinned to descendant head {PROOF_HEAD} for this package.",
        "An executable exact-identity crown gate was added. It cannot emit the final proved state unless exact source identity, body identity, playable IR, OneBody ABI, direct inline contact, visible consequence, no source mutation and identity preservation all pass. Marked fixtures cannot impersonate source proof.",
        "Exact Identity Gate v0.1 passed GitHub Actions. The exact-source seam in the adapter was then identified: the server already targets the frozen Untitled Field file, but the public repository intentionally does not contain those private donor bytes.",
        "A private exact-source intake was therefore added. The adapter can receive the donor through a private deployment input (JM_EXACT_BODY_B64) or local package file, hash the raw bytes before release to the UI resource, reject any mismatch, and report exactSourceVerified / observedSourceSha256 / sourceIntake. No frozen donor bytes were committed to the public repository.",
        "Exact Source Intake v0.1 also passed GitHub Actions, including syntax and mismatch-rejection checks. The public Cloudflare worker remains probe-only by design; private donor carriage and public safe-probe publication are kept separate.",
    ]),
    ("9. Repeatable operational loop", [
        "DEFINE BODY → DECLARE REQUIREMENTS → INSPECT HOST → COMPARE REQUIREMENTS/CAPABILITIES → SELECT OR COMPOSE ADAPTER → VERIFY SOURCE/IDENTITY → BIND MANIFESTATION → RUN → DIRECT CONTACT → FUNCTION TEST → IDENTITY TEST → TRACE/RECEIPT → STATUS.",
        "For future bodies, the architecture should reuse the recovered registries, adapter organs and proof receipts rather than inventing a new ‘inline runtime’ or bespoke portability philosophy each time.",
    ]),
    ("10. Proof ladder and claim ceiling", [
        "ARCHITECTURE / SPECIFICATION: COMPLETE for this closure scope.",
        "EXACT-IDENTITY GATE: BUILT + CI PASS.",
        "PRIVATE EXACT-SOURCE INTAKE: BUILT + CI PASS.",
        "CHATGPT INLINE HOST CLASS: DIRECT CONTACT PROVED historically by Third Ding.",
        "REFERENCE BODY CONVENTIONAL HOST CONTACT: PROVED for Untitled Field v0.9 within recovered phone/desktop evidence.",
        "EXACT SAME FROZEN REFERENCE BODY DIRECTLY CONTACTED INLINE THROUGH CHATGPT ADAPTER: OPEN — not yet claimed by this package.",
        "ARBITRARY-HOST / UNIVERSAL EXECUTION: OPEN — explicitly uncrowned.",
    ]),
    ("11. Canonical keeper", [
        "ONE BODY → KNOW ITS NEEDS → MEET THE HOST → ADAPT THE DIFFERENCE → VERIFY IDENTITY → MANIFEST → CONTACT → PROVE FUNCTION → PRESERVE IDENTITY.",
        "The body remains sovereign. The host is a carrier. The adapter reconciles difference. The manifestation is only the body’s current expression. Proof requires contact and trace.",
    ]),
]

EVIDENCE = [
    ["Reference", "Role / supported fact"],
    ["06_RESUME_CAPSULE_v0_9.md", "Untitled Field v0.9 frozen keeper; 27/27 phone-path QA, desktop/gesture pass, owner-device acceptance; exact SHA."],
    ["MCP_RUNTIME_PROOF_v0_1_1.json", "ChatGPT adapter runtime transport; bodyId, source SHA, playable IR, OneBody ABI; explicitly not direct ChatGPT contact proof."],
    ["BODY_POINTER.json", "Exact donor pointer; mutation_allowed=false; CARRIER_NOT_SOURCE_AUTHORITY."],
    ["JM_INLINE_CONTACT_FULL_CONVERGENCE_v1_0_2026-08-21", "Recovered organ stack and thin ChatGPT host-adapter correction."],
    ["JM INLINE CONTACT Third Ding v1.0", "Direct in-message interaction: READY/0 → DING 1/1 → DING 5/5; proves inline host-contact class."],
    ["PR #166 + descendant proof branch", "Non-destructive exact-identity proof rail, crown gate, private source intake and CI."],
]

STATUS = [
    ["Scope", "Status"],
    ["Architecture + ontology", "FROZEN · LOCKED · ANCHORED"],
    ["Adapter/proof-gate specification", "FROZEN · LOCKED · ANCHORED"],
    ["Exact-identity crown gate", "BUILT · CI PASS · anchored by descendant branch/PR"],
    ["Private exact-source intake", "BUILT · CI PASS · hash mismatch rejects"],
    ["Exact-body live ChatGPT contact", "OPEN EMPIRICAL GATE"],
    ["Universal arbitrary-host crown", "OPEN / NOT CLAIMED"],
]


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        if edge in kwargs:
            tag = 'w:' + edge
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, value in kwargs[edge].items():
                element.set(qn('w:' + key), str(value))


def add_docx_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("JM Universal Hosted-Body Architecture · v1.0 · Frozen architecture/proof-gate scope · Exact live-source inline contact remains open")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(110, 110, 110)


def build_docx():
    d = Document()
    sec = d.sections[0]
    sec.top_margin = Inches(0.62); sec.bottom_margin = Inches(0.62)
    sec.left_margin = Inches(0.68); sec.right_margin = Inches(0.68)
    add_docx_footer(sec)

    styles = d.styles
    styles['Normal'].font.name = 'Aptos'; styles['Normal'].font.size = Pt(10.2)
    styles['Heading 1'].font.name = 'Aptos Display'; styles['Heading 1'].font.size = Pt(18); styles['Heading 1'].font.bold = True
    styles['Heading 2'].font.name = 'Aptos Display'; styles['Heading 2'].font.size = Pt(13); styles['Heading 2'].font.bold = True

    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE); r.bold = True; r.font.size = Pt(23); r.font.name = 'Aptos Display'
    r.font.color.rgb = RGBColor(30, 38, 61)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SUBTITLE); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor(89, 67, 140)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(DATE); r.font.size = Pt(9); r.font.color.rgb = RGBColor(105, 105, 105)

    box = d.add_table(rows=4, cols=2); box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.style = 'Table Grid'
    metadata = [
        ("CANONICAL ROUTE", CANONICAL_ROUTE),
        ("CLOSURE STATUS", "FULL COMPLETE at architecture + adapter/proof-gate scope"),
        ("GOVERNANCE", "FROZEN · LOCKED · ANCHORED"),
        ("EMPIRICAL CEILING", "Exact frozen-body live ChatGPT contact remains OPEN"),
    ]
    for i,(k,v) in enumerate(metadata):
        box.cell(i,0).text = k; box.cell(i,1).text = v
        set_cell_shading(box.cell(i,0), '20273D'); set_cell_shading(box.cell(i,1), 'F4F1FA')
        for run in box.cell(i,0).paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255,255,255); run.bold = True; run.font.size = Pt(8.5)
        for run in box.cell(i,1).paragraphs[0].runs:
            run.font.size = Pt(9)
    d.add_paragraph()

    p = d.add_paragraph()
    r = p.add_run("EXECUTIVE KEEPER  "); r.bold=True; r.font.color.rgb = RGBColor(89,67,140)
    r = p.add_run("One body → know its needs → meet the host → adapt the difference → verify identity → manifest → contact → prove function → preserve identity."); r.bold=True

    d.add_heading('Architecture at a glance', level=1)
    t = d.add_table(rows=2, cols=5); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    labels = ['BODY','ADAPTER','HOST','MANIFESTATION','PROOF']
    values = ['identity + function','translate differences','execution/container','inline/panel/plugin/etc.','contact + consequence + receipt']
    for j in range(5):
        t.cell(0,j).text=labels[j]; t.cell(1,j).text=values[j]
        set_cell_shading(t.cell(0,j), '59438C'); set_cell_shading(t.cell(1,j), 'F3F0F9')
        for rr in t.cell(0,j).paragraphs[0].runs:
            rr.font.color.rgb=RGBColor(255,255,255); rr.bold=True; rr.font.size=Pt(8)
        for rr in t.cell(1,j).paragraphs[0].runs: rr.font.size=Pt(8.2)
        t.cell(0,j).vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; t.cell(1,j).vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

    d.add_heading('Plain-language separation', level=1)
    ct = d.add_table(rows=len(CORE), cols=2); ct.alignment=WD_TABLE_ALIGNMENT.CENTER; ct.style='Table Grid'
    for i,(k,v) in enumerate(CORE):
        ct.cell(i,0).text=k; ct.cell(i,1).text=v
        set_cell_shading(ct.cell(i,0),'E7E1F2')
        for rr in ct.cell(i,0).paragraphs[0].runs: rr.bold=True

    for title, paras in SECTIONS:
        d.add_heading(title, level=1)
        for text in paras:
            p=d.add_paragraph(style='Normal'); p.paragraph_format.space_after=Pt(5)
            r=p.add_run(text)
            if text in ("BODY ≠ HOST ≠ ADAPTER ≠ MANIFESTATION.",) or text.startswith("ONE BODY"):
                r.bold=True; r.font.color.rgb=RGBColor(55,47,91)

    d.add_heading('12. Exact reference identity', level=1)
    rt=d.add_table(rows=6, cols=2); rt.style='Table Grid'; rt.alignment=WD_TABLE_ALIGNMENT.CENTER
    refdata=[('Body ID',BODY_ID),('Frozen SHA-256',BODY_SHA),('Playable IR',PLAYABLE_IR),('OneBody ABI',ONEBODY_ABI),('Proof branch',PROOF_BRANCH),('Proof PR / head',f'{PROOF_PR} · {PROOF_HEAD}')]
    for i,(k,v) in enumerate(refdata):
        rt.cell(i,0).text=k; rt.cell(i,1).text=v; set_cell_shading(rt.cell(i,0),'E7E1F2')
        for rr in rt.cell(i,0).paragraphs[0].runs: rr.bold=True
        for rr in rt.cell(i,1).paragraphs[0].runs: rr.font.name='Consolas'; rr.font.size=Pt(8.2)

    d.add_heading('13. Status matrix', level=1)
    st=d.add_table(rows=len(STATUS), cols=2); st.style='Table Grid'; st.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,row in enumerate(STATUS):
        for j,val in enumerate(row): st.cell(i,j).text=val
        if i==0:
            for c in st.rows[i].cells:
                set_cell_shading(c,'20273D')
                for rr in c.paragraphs[0].runs: rr.font.color.rgb=RGBColor(255,255,255); rr.bold=True
        else:
            if 'OPEN' in row[1]: set_cell_shading(st.cell(i,1),'FFF1E3')
            elif 'PASS' in row[1] or 'FROZEN' in row[1]: set_cell_shading(st.cell(i,1),'EAF6EC')

    d.add_heading('14. Evidence basis', level=1)
    et=d.add_table(rows=len(EVIDENCE), cols=2); et.style='Table Grid'; et.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,row in enumerate(EVIDENCE):
        for j,val in enumerate(row): et.cell(i,j).text=val
        if i==0:
            for c in et.rows[i].cells:
                set_cell_shading(c,'20273D')
                for rr in c.paragraphs[0].runs: rr.font.color.rgb=RGBColor(255,255,255); rr.bold=True

    d.add_heading('15. Freeze / Lock / Anchor receipt', level=1)
    receipt = [
        "FREEZE — This document fixes the architecture, terminology, proof hierarchy, recovered organ routing and 28 August 2026 proof-gate implementation at v1.0. Frozen historical donors and Third Ding evidence remain unchanged.",
        "LOCK — BODY ≠ HOST ≠ ADAPTER ≠ MANIFESTATION; host carriage never grants source authority; source identity must be verified before stronger same-body claims; contact and consequence are required for direct-function proof.",
        "ANCHOR — Canonical route: " + CANONICAL_ROUTE,
        "OPEN GATE — The exact frozen Untitled Field v0.9 donor has not yet been directly contacted inline through the live private exact-source ChatGPT deployment. That empirical join remains open and cannot be inherited from the safe probe.",
        "NO UNIVERSAL CROWN — This closure does not claim arbitrary-host execution, universal availability, or identical implementation on every host.",
    ]
    for x in receipt:
        p=d.add_paragraph(style='Normal'); p.style=d.styles['Normal']; p.paragraph_format.left_indent=Inches(0.18)
        r=p.add_run(x); r.bold=True if x.startswith(('FREEZE','LOCK','ANCHOR','OPEN GATE','NO UNIVERSAL')) else False

    d.add_heading('16. Storage-clean package rule', level=1)
    d.add_paragraph("Canonical storage set is exactly one DOCX + one PDF. The governance receipt is embedded in-body. No loose receipt and no ZIP are required for the canonical set.")
    d.add_paragraph("This package may be superseded only by an explicitly declared descendant after new evidence or material architecture change. The frozen v1.0 closure itself is not silently rewritten.")

    d.save(DOCX)


def pdf_header_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont('Helvetica', 7.4)
    canvas.setFillColor(colors.HexColor('#777777'))
    canvas.drawCentredString(A4[0]/2, 9*mm, f"JM Universal Hosted-Body Architecture · v1.0 · Page {doc.page}")
    canvas.restoreState()


def build_pdf():
    styles=getSampleStyleSheet()
    title=ParagraphStyle('TitleX', parent=styles['Title'], fontName='Helvetica-Bold', fontSize=22, leading=24, textColor=colors.HexColor('#20273D'), alignment=TA_CENTER, spaceAfter=6)
    sub=ParagraphStyle('SubX', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=10.7, leading=14, textColor=colors.HexColor('#59438C'), alignment=TA_CENTER, spaceAfter=5)
    body=ParagraphStyle('BodyX', parent=styles['BodyText'], fontName='Helvetica', fontSize=9.3, leading=13.4, textColor=colors.HexColor('#242424'), spaceAfter=6)
    h1=ParagraphStyle('H1X', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=14.2, leading=17, textColor=colors.HexColor('#20273D'), spaceBefore=10, spaceAfter=6)
    keeper=ParagraphStyle('Keeper', parent=body, fontName='Helvetica-Bold', fontSize=10.2, leading=14, textColor=colors.HexColor('#443467'), borderColor=colors.HexColor('#CFC4E5'), borderWidth=.7, borderPadding=8, backColor=colors.HexColor('#F6F3FA'), spaceBefore=7, spaceAfter=10)
    small=ParagraphStyle('Small', parent=body, fontSize=7.8, leading=10.2)

    doc=SimpleDocTemplate(str(PDF), pagesize=A4, rightMargin=16*mm,leftMargin=16*mm,topMargin=15*mm,bottomMargin=17*mm,title=TITLE,author='JM / JMISJUSTME')
    story=[]
    story += [Paragraph(TITLE,title), Paragraph(SUBTITLE,sub), Paragraph(DATE, ParagraphStyle('date',parent=small,alignment=TA_CENTER,textColor=colors.HexColor('#777777'))), Spacer(1,4*mm)]
    meta=[["CANONICAL ROUTE",CANONICAL_ROUTE],["CLOSURE STATUS","FULL COMPLETE at architecture + adapter/proof-gate scope"],["GOVERNANCE","FROZEN · LOCKED · ANCHORED"],["EMPIRICAL CEILING","Exact frozen-body live ChatGPT contact remains OPEN"]]
    mt=Table(meta,colWidths=[40*mm,138*mm])
    mt.setStyle(TableStyle([('BACKGROUND',(0,0),(0,-1),colors.HexColor('#20273D')),('TEXTCOLOR',(0,0),(0,-1),colors.white),('BACKGROUND',(1,0),(1,-1),colors.HexColor('#F4F1FA')),('FONTNAME',(0,0),(0,-1),'Helvetica-Bold'),('FONTNAME',(1,0),(1,-1),'Helvetica'),('FONTSIZE',(0,0),(-1,-1),7.8),('LEADING',(0,0),(-1,-1),10),('VALIGN',(0,0),(-1,-1),'MIDDLE'),('GRID',(0,0),(-1,-1),.3,colors.HexColor('#D8D3E2')),('LEFTPADDING',(0,0),(-1,-1),5),('RIGHTPADDING',(0,0),(-1,-1),5),('TOPPADDING',(0,0),(-1,-1),5),('BOTTOMPADDING',(0,0),(-1,-1),5)]))
    story += [mt, Spacer(1,4*mm), Paragraph("ONE BODY → KNOW ITS NEEDS → MEET THE HOST → ADAPT THE DIFFERENCE → VERIFY IDENTITY → MANIFEST → CONTACT → PROVE FUNCTION → PRESERVE IDENTITY.",keeper)]

    story.append(Paragraph('Architecture at a glance',h1))
    row1=[Paragraph(x,ParagraphStyle('th',parent=small,fontName='Helvetica-Bold',textColor=colors.white,alignment=TA_CENTER)) for x in ['BODY','ADAPTER','HOST','MANIFESTATION','PROOF']]
    row2=[Paragraph(x,ParagraphStyle('td',parent=small,alignment=TA_CENTER)) for x in ['identity + function','translate differences','execution/container','inline/panel/plugin/etc.','contact + consequence + receipt']]
    at=Table([row1,row2],colWidths=[35.6*mm]*5)
    at.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),colors.HexColor('#59438C')),('BACKGROUND',(0,1),(-1,1),colors.HexColor('#F3F0F9')),('GRID',(0,0),(-1,-1),.35,colors.HexColor('#D8D3E2')),('VALIGN',(0,0),(-1,-1),'MIDDLE'),('TOPPADDING',(0,0),(-1,-1),6),('BOTTOMPADDING',(0,0),(-1,-1),6)]))
    story += [at,Spacer(1,3*mm)]

    story.append(Paragraph('Plain-language separation',h1))
    ct=Table([[Paragraph(k,ParagraphStyle('k',parent=small,fontName='Helvetica-Bold')),Paragraph(v,small)] for k,v in CORE],colWidths=[34*mm,144*mm])
    ct.setStyle(TableStyle([('BACKGROUND',(0,0),(0,-1),colors.HexColor('#E7E1F2')),('GRID',(0,0),(-1,-1),.35,colors.HexColor('#D8D3E2')),('VALIGN',(0,0),(-1,-1),'TOP'),('LEFTPADDING',(0,0),(-1,-1),5),('RIGHTPADDING',(0,0),(-1,-1),5),('TOPPADDING',(0,0),(-1,-1),5),('BOTTOMPADDING',(0,0),(-1,-1),5)]))
    story += [ct,Spacer(1,2*mm)]

    for heading,paras in SECTIONS:
        story.append(Paragraph(heading,h1))
        for txt in paras:
            st=keeper if txt.startswith('BODY ≠') or txt.startswith('ONE BODY') else body
            story.append(Paragraph(txt.replace('→','&#8594;').replace('↔','&#8596;').replace('≠','&#8800;'),st))

    story.append(Paragraph('12. Exact reference identity',h1))
    refdata=[('Body ID',BODY_ID),('Frozen SHA-256',BODY_SHA),('Playable IR',PLAYABLE_IR),('OneBody ABI',ONEBODY_ABI),('Proof branch',PROOF_BRANCH),('Proof PR / head',f'{PROOF_PR} · {PROOF_HEAD}')]
    rt=Table([[Paragraph(k,ParagraphStyle('rk',parent=small,fontName='Helvetica-Bold')),Paragraph(v,ParagraphStyle('rv',parent=small,fontName='Courier',fontSize=6.9,leading=9))] for k,v in refdata],colWidths=[38*mm,140*mm])
    rt.setStyle(TableStyle([('BACKGROUND',(0,0),(0,-1),colors.HexColor('#E7E1F2')),('GRID',(0,0),(-1,-1),.35,colors.HexColor('#D8D3E2')),('VALIGN',(0,0),(-1,-1),'TOP'),('LEFTPADDING',(0,0),(-1,-1),5),('RIGHTPADDING',(0,0),(-1,-1),5),('TOPPADDING',(0,0),(-1,-1),5),('BOTTOMPADDING',(0,0),(-1,-1),5)]))
    story.append(rt)

    story.append(Paragraph('13. Status matrix',h1))
    srows=[[Paragraph(c,ParagraphStyle('sth' if i==0 else 'std',parent=small,fontName='Helvetica-Bold' if i==0 else 'Helvetica',textColor=colors.white if i==0 else colors.HexColor('#242424'))) for c in row] for i,row in enumerate(STATUS)]
    st=Table(srows,colWidths=[86*mm,92*mm])
    style=[('BACKGROUND',(0,0),(-1,0),colors.HexColor('#20273D')),('GRID',(0,0),(-1,-1),.35,colors.HexColor('#D8D3E2')),('VALIGN',(0,0),(-1,-1),'TOP'),('LEFTPADDING',(0,0),(-1,-1),5),('RIGHTPADDING',(0,0),(-1,-1),5),('TOPPADDING',(0,0),(-1,-1),5),('BOTTOMPADDING',(0,0),(-1,-1),5)]
    for i,row in enumerate(STATUS[1:],1):
        style.append(('BACKGROUND',(1,i),(1,i),colors.HexColor('#FFF1E3') if 'OPEN' in row[1] else colors.HexColor('#EAF6EC')))
    st.setStyle(TableStyle(style)); story.append(st)

    story.append(Paragraph('14. Evidence basis',h1))
    erows=[[Paragraph(c,ParagraphStyle('eth' if i==0 else 'etd',parent=small,fontName='Helvetica-Bold' if i==0 else 'Helvetica',textColor=colors.white if i==0 else colors.HexColor('#242424'))) for c in row] for i,row in enumerate(EVIDENCE)]
    et=Table(erows,colWidths=[58*mm,120*mm],repeatRows=1)
    et.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,0),colors.HexColor('#20273D')),('GRID',(0,0),(-1,-1),.35,colors.HexColor('#D8D3E2')),('VALIGN',(0,0),(-1,-1),'TOP'),('LEFTPADDING',(0,0),(-1,-1),5),('RIGHTPADDING',(0,0),(-1,-1),5),('TOPPADDING',(0,0),(-1,-1),5),('BOTTOMPADDING',(0,0),(-1,-1),5)])); story.append(et)

    story.append(Paragraph('15. Freeze / Lock / Anchor receipt',h1))
    rec=[
        '<b>FREEZE —</b> Architecture, terminology, proof hierarchy, recovered organ routing and 28 August 2026 proof-gate implementation are fixed at v1.0. Historical frozen donors and Third Ding evidence remain unchanged.',
        '<b>LOCK —</b> BODY ≠ HOST ≠ ADAPTER ≠ MANIFESTATION; host carriage never grants source authority; exact identity must be verified before stronger same-body claims; direct contact + visible consequence are required for direct-function proof.',
        '<b>ANCHOR —</b> Canonical route: '+CANONICAL_ROUTE,
        '<b>OPEN GATE —</b> Exact frozen Untitled Field v0.9 has not yet been directly contacted inline through the live private exact-source ChatGPT deployment. The safe probe cannot substitute for this join.',
        '<b>NO UNIVERSAL CROWN —</b> Arbitrary-host execution, universal availability and identical implementation on every host are not claimed.'
    ]
    for x in rec: story.append(Paragraph(x,keeper))

    story.append(Paragraph('16. Storage-clean package rule',h1))
    story.append(Paragraph('Canonical storage set is exactly one DOCX + one PDF. Governance/status receipt is embedded in-body. No loose receipt and no ZIP are required for the canonical set.',body))
    story.append(Paragraph('This package may be superseded only by an explicitly declared descendant after new evidence or material architecture change. The frozen v1.0 closure itself is not silently rewritten.',body))

    doc.build(story,onFirstPage=pdf_header_footer,onLaterPages=pdf_header_footer)


def main():
    build_docx(); build_pdf()
    for p in (DOCX,PDF):
        data=p.read_bytes()
        print(f"{p.name}\t{len(data)} bytes\tsha256={sha256(data).hexdigest()}")

if __name__ == '__main__':
    main()
